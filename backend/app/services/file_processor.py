import os
import hashlib
import logging
import base64
import io
import tempfile
import re
from typing import List, Dict, Optional, Tuple, Callable, Any
from concurrent.futures import ThreadPoolExecutor, as_completed
import fitz  # PyMuPDF
from docx import Document
from openai import OpenAI
import voyageai
import google.generativeai as genai
from PIL import Image
from app.config import get_settings

settings = get_settings()
logger = logging.getLogger(__name__)

# Initialize DeepSeek client (OpenAI-compatible) for chat
deepseek_client = None
if settings.deepseek_api_key:
    deepseek_client = OpenAI(
        api_key=settings.deepseek_api_key,
        base_url="https://api.deepseek.com/v1"
    )

# Initialize Voyage AI client for embeddings
voyage_client = None
if settings.voyage_api_key:
    voyage_client = voyageai.Client(api_key=settings.voyage_api_key)
    logger.info("Voyage AI configured for embeddings (voyage-3-large, 1024 dims)")

# Initialize Gemini client for vision/image processing
gemini_vision_client = None
GEMINI_VISION_MODEL = settings.gemini_model or "gemini-2.5-flash"
if settings.gemini_api_key:
    genai.configure(api_key=settings.gemini_api_key)
    gemini_vision_client = genai.GenerativeModel(GEMINI_VISION_MODEL)
    logger.info(f"{GEMINI_VISION_MODEL} configured for vision/image processing")

# Embedding configuration
EMBEDDING_MODEL = "voyage-3-large"  # Best quality: 1024 dimensions
EMBEDDING_DIMENSIONS = 1024

# Image file extensions
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff', '.tif'}
IMAGE_PROCESSING_MAX_WORKERS = 5
IMAGE_PROCESSING_BATCH_SIZE = 50


def _compact_text(value: str, max_len: int = 180) -> str:
    if not value:
        return ""
    text = re.sub(r"\s+", " ", value).strip(" -•\t")
    if len(text) <= max_len:
        return text
    return text[:max_len].rstrip(" ,.;:") + "..."


def _extract_indicator_hint(analysis: str) -> str:
    """
    Extract a short, user-facing observation from Gemini output.
    This keeps indicators grounded in what Gemini actually saw.
    """
    if not analysis:
        return ""

    normalized = analysis.replace("\r\n", "\n").replace("\r", "\n")
    headings = [
        "KEY INSIGHTS:",
        "DETAILED DESCRIPTION:",
        "VISUAL LAYOUT:",
        "IMAGE TYPE:",
    ]

    for heading in headings:
        idx = normalized.find(heading)
        if idx == -1:
            continue
        tail = normalized[idx + len(heading):].strip()
        if not tail:
            continue
        candidate = tail.split("\n\n", 1)[0].strip()
        if "\n" in candidate:
            candidate = candidate.split("\n", 1)[0].strip()
        candidate = _compact_text(candidate)
        if candidate:
            return candidate

    first_nonempty_line = ""
    for line in normalized.split("\n"):
        cleaned = _compact_text(line)
        if cleaned:
            first_nonempty_line = cleaned
            break
    return first_nonempty_line


def _emit_progress(
    progress_callback: Optional[Callable[..., Any]],
    stage: str,
    label: str,
    current: int,
    total: int,
    meta: Optional[Dict] = None,
    force: bool = False,
) -> None:
    if not progress_callback:
        return
    try:
        progress_callback(
            stage=stage,
            label=label,
            current=current,
            total=total,
            meta=meta or {},
            force=force,
        )
    except Exception as e:
        logger.debug(f"Progress callback failed (non-critical): {e}")


def is_image_file(file_path: str, mime_type: str) -> bool:
    """Check if file is an image"""
    file_ext = os.path.splitext(file_path)[1].lower()
    return (
        file_ext in IMAGE_EXTENSIONS or 
        mime_type.startswith('image/')
    )


def process_image_with_vision(file_path: str, filename: str = "") -> Dict:
    """
    Process image using Gemini Vision to extract text and describe content.

    Returns:
        Dict with extracted text and metadata
    """
    logger.info(f"🖼️  Processing image with {GEMINI_VISION_MODEL}: {filename}")

    if not gemini_vision_client:
        logger.error("❌ Gemini API not configured - image processing requires GEMINI_API_KEY")
        raise Exception("Gemini API not configured - image processing requires GEMINI_API_KEY")

    try:
        logger.debug(f"  📖 Reading image file: {file_path}")
        img = Image.open(file_path)

        prompt = f"""Analyze this image ({filename if filename else 'image'}) with MAXIMUM detail. Extract absolutely everything visible.

Format your response EXACTLY as:

IMAGE TYPE:
[One of: Photograph, Screenshot, Chart/Graph, Diagram/Flowchart, Table, Icon/Logo, Illustration, UI/Landing Page, Infographic, Map, Document Scan, Other]

TEXT CONTENT (OCR):
[Extract ALL text exactly as it appears - every word, number, label, caption, watermark, button text, menu item. If no text, write "No text detected"]

VISUAL LAYOUT:
[Describe the complete layout: positioning of elements (top-left, center, bottom-right), spacing, alignment, grid structure, columns, sections, hierarchy of visual elements]

COLORS & DESIGN:
[List ALL colors used with hex codes where possible. Describe: background color, text colors, accent colors, gradients, shadows, borders, color scheme (dark/light/colorful), contrast]

TYPOGRAPHY:
[Describe fonts: serif/sans-serif, bold/regular/light, sizes (heading vs body), font colors, text alignment, any decorative text]

DETAILED DESCRIPTION:
[Comprehensive description of EVERYTHING in the image: objects, people (appearance, clothing, expressions, positions), scenes, backgrounds, foregrounds, textures, patterns, shapes, icons, buttons, UI elements, branding]

DATA & NUMBERS:
[For charts/graphs: axis labels, all data points, values, trends, percentages, comparisons. For tables: all rows and columns. For any numbers/statistics: extract them all with context]

DESIGN ANALYSIS:
[If UI/web/app screenshot: framework hints (Material, Bootstrap, Tailwind), responsive design, navigation structure, CTA placement, hero section, footer. If marketing: target audience, messaging strategy, visual hierarchy]

KEY INSIGHTS:
[Important takeaways, notable patterns, relationships between elements, anything remarkable or unusual]"""

        logger.info(f"  🤖 Calling Gemini Vision API ({GEMINI_VISION_MODEL}) for {filename}...")
        response = gemini_vision_client.generate_content([prompt, img])
        analysis = response.text
        logger.info(f"  ✅ Gemini Vision analysis complete for {filename}")

        text_content = analysis
        if "TEXT CONTENT" in analysis:
            # Accept either new or older heading variants from model output.
            split_marker = "VISUAL LAYOUT:" if "VISUAL LAYOUT:" in analysis else "VISUAL DESCRIPTION:"
            text_part = analysis.split(split_marker, 1)[0]
            text_part = text_part.replace("TEXT CONTENT (OCR):", "").replace("TEXT CONTENT:", "").strip()
            text_content = f"{text_part}\n\n{analysis}" if text_part else analysis

        word_count = len(text_content.split())
        logger.debug(f"  📊 Extracted {word_count} words from image")
        indicator_hint = _extract_indicator_hint(analysis)

        return {
            "text": text_content,
            "metadata": {
                "type": "image",
                "word_count": word_count,
                "char_count": len(text_content),
                "vision_model": GEMINI_VISION_MODEL,
                "indicator_hint": indicator_hint,
            }
        }

    except Exception as e:
        logger.error(f"❌ Image processing failed for {filename}: {str(e)}", exc_info=True)
        raise Exception(f"Failed to process image: {str(e)}")


def extract_text_from_file(
    file_path: str,
    mime_type: str,
    progress_callback: Optional[Callable[..., Any]] = None
) -> Dict:
    """Extract text from various file types - handles text, PDFs, images, and more"""
    text = ""
    metadata = {}
    
    try:
        # Get file extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Handle IMAGE files with Kimi Vision
        if is_image_file(file_path, mime_type):
            logger.info(f"Processing image file with Gemini Vision: {file_path}")
            _emit_progress(
                progress_callback,
                "image_ocr",
                "Opening your image and checking what is in it...",
                0,
                1,
                {"file_type": "image"},
                True,
            )
            result = process_image_with_vision(file_path, os.path.basename(file_path))
            hint = result.get("metadata", {}).get("indicator_hint") or ""
            label = "Image review complete."
            if hint:
                label = f"I can see: {hint}"
            _emit_progress(
                progress_callback,
                "image_ocr",
                label,
                1,
                1,
                {"file_type": "image", "vision_hint": hint},
                True,
            )
            return result
        
        # Handle PDF files with PyMuPDF (extracts text AND images)
        if mime_type == "application/pdf" or file_ext == ".pdf":
            logger.info(f"📄 Processing PDF with PyMuPDF: {file_path}")
            doc = fitz.open(file_path)
            text = ""
            images_processed = 0
            page_count = len(doc)

            # PHASE 1: Extract all text and collect all images (fast, local)
            page_texts = {}
            image_tasks = []  # List of (page_num, img_index, temp_path, label)

            for page_num in range(page_count):
                page = doc[page_num]

                # Extract text from page
                page_text = page.get_text()
                page_texts[page_num] = page_text

                # Extract and save all images to temp files
                image_list = page.get_images()
                if image_list:
                    logger.info(f"  🖼️  Found {len(image_list)} images on page {page_num + 1}")

                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]

                            # Save image temporarily
                            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{image_ext}') as temp_img:
                                temp_img.write(image_bytes)
                                temp_img_path = temp_img.name

                            label = f"PDF_Page{page_num + 1}_Image{img_index + 1}"
                            image_tasks.append((page_num, img_index, temp_img_path, label))

                        except Exception as e:
                            logger.warning(f"    ⚠️  Failed to extract image {img_index + 1} on page {page_num + 1}: {str(e)}")
                            continue

            doc.close()

            total_images = len(image_tasks)
            logger.info(
                f"  📊 Extracted {total_images} images from {page_count} pages "
                f"- processing in batches of {IMAGE_PROCESSING_BATCH_SIZE} "
                f"with up to {IMAGE_PROCESSING_MAX_WORKERS} workers"
            )
            _emit_progress(
                progress_callback,
                "image_ocr",
                (
                    f"I found {total_images} visuals. Reviewing them one by one..."
                    if total_images else
                    "No visuals found in this PDF."
                ),
                0,
                total_images,
                {"file_type": "pdf", "page_count": page_count},
                True,
            )

            # PHASE 2: Process ALL images in parallel (the speed boost)
            image_results = {}  # (page_num, img_index) -> result text

            if image_tasks:
                images_completed = 0
                def _process_single_image(task):
                    """Worker function for parallel image processing"""
                    page_num, img_index, temp_path, label = task
                    try:
                        result = process_image_with_vision(temp_path, label)
                        hint = result.get("metadata", {}).get("indicator_hint", "")
                        return (page_num, img_index, result['text'], hint, temp_path, None)
                    except Exception as e:
                        return (page_num, img_index, None, "", temp_path, str(e))

                total_batches = (total_images + IMAGE_PROCESSING_BATCH_SIZE - 1) // IMAGE_PROCESSING_BATCH_SIZE
                for batch_index in range(total_batches):
                    start = batch_index * IMAGE_PROCESSING_BATCH_SIZE
                    end = min(start + IMAGE_PROCESSING_BATCH_SIZE, total_images)
                    batch_tasks = image_tasks[start:end]
                    workers = min(IMAGE_PROCESSING_MAX_WORKERS, len(batch_tasks))

                    logger.info(
                        f"  🚚 Processing image batch {batch_index + 1}/{total_batches} "
                        f"({len(batch_tasks)} images, workers={workers})"
                    )

                    with ThreadPoolExecutor(max_workers=workers) as executor:
                        futures = {executor.submit(_process_single_image, task): task for task in batch_tasks}

                        for future in as_completed(futures):
                            page_num, img_index, result_text, hint, temp_path, error = future.result()
                            images_completed += 1

                            # Clean up temp file
                            try:
                                os.unlink(temp_path)
                            except OSError:
                                pass

                            if error:
                                logger.warning(f"    ⚠️  Failed to process image {img_index + 1} on page {page_num + 1}: {error}")
                            else:
                                image_results[(page_num, img_index)] = result_text
                                images_processed += 1
                                logger.info(f"    ✅ Image {img_index + 1} on page {page_num + 1} done ({images_processed}/{total_images})")
                            progress_label = f"Reviewing visuals ({images_completed}/{total_images})"
                            hint_text = _compact_text(hint, max_len=140)
                            if hint_text:
                                progress_label = f"Image {images_completed}/{total_images}: {hint_text}"
                            _emit_progress(
                                progress_callback,
                                "image_ocr",
                                progress_label,
                                images_completed,
                                total_images,
                                {
                                    "file_type": "pdf",
                                    "page": page_num + 1,
                                    "image_index": img_index + 1,
                                    "total_images": total_images,
                                    "vision_hint": hint_text,
                                    "batch_index": batch_index + 1,
                                    "batch_total": total_batches,
                                },
                            )

            # PHASE 3: Assemble final text in correct page/image order + build per-page structure
            pages_structured = []
            for page_num in range(page_count):
                text += f"\n{'='*60}\n=== PAGE {page_num + 1} of {page_count} ===\n{'='*60}\n{page_texts[page_num]}"

                # Add image results for this page in order
                page_images_sorted = sorted(
                    [(idx, txt) for (pn, idx), txt in image_results.items() if pn == page_num],
                    key=lambda x: x[0]
                )
                for img_index, img_text in page_images_sorted:
                    text += f"\n{'─'*40}\n[IMAGE {img_index + 1} - Page {page_num + 1}]\n{img_text}\n{'─'*40}\n"

                # Build structured per-page data for spatial summary
                pages_structured.append({
                    "page_num": page_num + 1,
                    "text": page_texts[page_num],
                    "images": [
                        {"label": f"Image {idx + 1}", "text": img_text}
                        for idx, img_text in page_images_sorted
                    ]
                })

            metadata["page_count"] = page_count
            metadata["images_extracted"] = images_processed
            metadata["pages"] = pages_structured  # Per-page structure for spatial summary
            logger.info(f"  ✅ PDF processed: {page_count} pages, {images_processed} images extracted (PARALLEL)")
            
        # Handle Word documents
        elif mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or file_ext in [".doc", ".docx"]:
            from docx.oxml.ns import qn
            doc = Document(file_path)

            # Build sections based on headings — each heading = one section
            sections = []
            current_section = {"heading": "Introduction", "paragraphs": [], "images": []}
            image_tasks = []  # (section_ref, img_bytes, img_ext)

            for para in doc.paragraphs:
                style_name = para.style.name if para.style else ""
                is_heading = "Heading" in style_name or style_name in ["Title", "Subtitle"]

                if is_heading and para.text.strip():
                    # Save current section, start new one
                    if current_section["paragraphs"] or current_section["images"]:
                        sections.append(current_section)
                    current_section = {"heading": para.text.strip(), "paragraphs": [], "images": []}
                else:
                    if para.text.strip():
                        current_section["paragraphs"].append(para.text.strip())

                # Collect inline images from this paragraph
                drawings = para._element.findall('.//' + qn('w:drawing'))
                for drawing in drawings:
                    blips = drawing.findall('.//' + qn('a:blip'))
                    for blip in blips:
                        r_embed = blip.get(qn('r:embed'))
                        if r_embed and r_embed in para.part.rels:
                            try:
                                img_part = para.part.rels[r_embed].target_part
                                img_ext = img_part.content_type.split('/')[-1]
                                if img_ext == 'jpeg':
                                    img_ext = 'jpg'
                                image_tasks.append((current_section, img_part.blob, img_ext))
                            except Exception as e:
                                logger.warning(f"  ⚠️  Failed to collect DOCX image: {e}")

            # Don't forget last section
            if current_section["paragraphs"] or current_section["images"]:
                sections.append(current_section)

            # If no headings found, treat whole doc as one section
            if not sections:
                all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                sections = [{"heading": os.path.basename(file_path), "paragraphs": [all_text], "images": []}]

            total_images = len(image_tasks)
            logger.info(
                f"  📊 DOCX: {len(sections)} sections, {total_images} images "
                f"- processing in batches of {IMAGE_PROCESSING_BATCH_SIZE} "
                f"with up to {IMAGE_PROCESSING_MAX_WORKERS} workers"
            )
            _emit_progress(
                progress_callback,
                "image_ocr",
                (
                    f"I found {total_images} visuals in this document. Reviewing each one..."
                    if total_images else
                    "No visuals found in this document."
                ),
                0,
                total_images,
                {"file_type": "docx", "section_count": len(sections)},
                True,
            )

            # Process all images in parallel, attach to their section
            if image_tasks:
                images_completed = 0
                def _process_docx_image(task):
                    section_ref, img_bytes, img_ext = task
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{img_ext}') as tmp:
                            tmp.write(img_bytes)
                            tmp_path = tmp.name
                        label = f"Image {len(section_ref['images']) + 1} in {section_ref['heading'][:30]}"
                        result = process_image_with_vision(tmp_path, label)
                        hint = result.get("metadata", {}).get("indicator_hint", "")
                        try:
                            os.unlink(tmp_path)
                        except OSError:
                            pass
                        return section_ref, result['text'], hint, None
                    except Exception as e:
                        return section_ref, None, "", str(e)

                total_batches = (total_images + IMAGE_PROCESSING_BATCH_SIZE - 1) // IMAGE_PROCESSING_BATCH_SIZE
                for batch_index in range(total_batches):
                    start = batch_index * IMAGE_PROCESSING_BATCH_SIZE
                    end = min(start + IMAGE_PROCESSING_BATCH_SIZE, total_images)
                    batch_tasks = image_tasks[start:end]
                    workers = min(IMAGE_PROCESSING_MAX_WORKERS, len(batch_tasks))

                    logger.info(
                        f"  🚚 DOCX image batch {batch_index + 1}/{total_batches} "
                        f"({len(batch_tasks)} images, workers={workers})"
                    )

                    with ThreadPoolExecutor(max_workers=workers) as executor:
                        futures = {executor.submit(_process_docx_image, task): task for task in batch_tasks}
                        for future in as_completed(futures):
                            section_ref, result_text, hint, error = future.result()
                            images_completed += 1
                            if error:
                                logger.warning(f"  ⚠️  DOCX image failed: {error}")
                            elif result_text:
                                section_ref["images"].append({
                                    "label": f"Image {len(section_ref['images']) + 1}",
                                    "text": result_text
                                })
                            progress_label = f"Reviewing visuals ({images_completed}/{total_images})"
                            hint_text = _compact_text(hint, max_len=140)
                            if hint_text:
                                progress_label = f"Image {images_completed}/{total_images}: {hint_text}"
                            _emit_progress(
                                progress_callback,
                                "image_ocr",
                                progress_label,
                                images_completed,
                                total_images,
                                {
                                    "file_type": "docx",
                                    "section": section_ref.get("heading"),
                                    "total_images": total_images,
                                    "vision_hint": hint_text,
                                    "batch_index": batch_index + 1,
                                    "batch_total": total_batches,
                                },
                            )

            # Assemble merged text + per-section structured data (same shape as PDF pages)
            sections_structured = []
            for s_idx, section in enumerate(sections):
                section_header = f"\n{'='*60}\n=== SECTION {s_idx + 1}: {section['heading']} ===\n{'='*60}\n"
                section_body = "\n".join(section["paragraphs"])
                text += section_header + section_body

                for img in section["images"]:
                    text += f"\n{'─'*40}\n[{img['label']}]\n{img['text']}\n{'─'*40}\n"

                sections_structured.append({
                    "page_num": s_idx + 1,
                    "text": section_body,
                    "images": section["images"],
                    "heading": section["heading"]
                })

            images_extracted = sum(len(s["images"]) for s in sections)
            metadata["page_count"] = len(sections)
            metadata["images_extracted"] = images_extracted
            metadata["pages"] = sections_structured  # Same structure as PDF for spatial summary
            logger.info(f"  ✅ DOCX processed: {len(sections)} sections, {images_extracted} images extracted")
            
        # Handle CSV files
        elif mime_type == "text/csv" or file_ext == ".csv":
            import csv
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
                text = "\n".join([", ".join(row) for row in rows])
                
        # Handle text-based code files and other text formats
        elif mime_type.startswith("text/") or file_ext in [
            ".txt", ".md", ".markdown", ".py", ".js", ".jsx", ".ts", ".tsx", 
            ".html", ".htm", ".css", ".json", ".xml", ".yaml", ".yml", ".sh", ".bash",
            ".c", ".cpp", ".cc", ".cxx", ".h", ".hpp", ".java", ".go", ".rs", ".rb", ".php",
            ".swift", ".kt", ".scala", ".sql", ".r", ".m", ".pl", ".pm",
            ".lua", ".vim", ".conf", ".config", ".log", ".ini", ".toml",
            ".dockerfile", ".env", ".gitignore", ".readme", ".mdx", ".tex", ".rst",
            ".scss", ".sass", ".less", ".styl", ".vue", ".svelte",
            ".dart", ".elm", ".ex", ".exs", ".clj", ".cljs", ".hs", ".ml", ".fs",
            ".ps1", ".psm1", ".bat", ".cmd"
        ]:
            # Try multiple encodings for text files
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252', 'utf-16']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        text = f.read()
                        break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            if not text:
                raise Exception("File appears to be binary or uses unsupported encoding")
        
        # If no specific handler matched, try to read as text (generic fallback)
        else:
            # Try to read as text with multiple encodings
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        text = f.read()
                        
                        # Basic binary detection: check for null bytes or too many non-printable chars
                        if len(text) > 0:
                            null_bytes = text.count('\x00')
                            non_printable = sum(1 for c in text[:min(1000, len(text))] 
                                              if ord(c) < 32 and c not in '\n\r\t\f\v')
                            
                            # If more than 5% null bytes or 10% non-printable, likely binary
                            if null_bytes > len(text) * 0.05 or non_printable > len(text[:min(1000, len(text))]) * 0.1:
                                text = ""
                                continue
                        break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            if not text:
                # Binary file - can't extract text
                raise Exception(f"File type '{mime_type}' ({file_ext}) appears to be binary and cannot be processed as text. Please upload a text-based file.")
                
        metadata["word_count"] = len(text.split()) if text else 0
        metadata["char_count"] = len(text) if text else 0
        
    except Exception as e:
        raise Exception(f"Failed to extract text: {str(e)}")
    
    return {"text": text, "metadata": metadata}


def chunk_text(text: str, chunk_size: int = 150, overlap: int = 50) -> List[Dict]:
    """
    Split text into overlapping chunks for maximum precision semantic search.
    
    Default: 150 words per chunk with 50-word overlap for high-quality retrieval.
    """
    if not text or len(text.strip()) == 0:
        return []
    
    words = text.split()
    chunks = []
    
    i = 0
    chunk_index = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        chunk_content = " ".join(chunk_words)
        
        chunks.append({
            "content": chunk_content,
            "chunk_index": chunk_index,
            "start_offset": i,
            "end_offset": min(i + chunk_size, len(words)),
            "token_count": len(chunk_words)
        })
        
        i += chunk_size - overlap
        chunk_index += 1
        
    return chunks


def generate_embedding(text: str) -> Optional[List[float]]:
    """
    Generate embedding using Voyage AI voyage-3-large model.

    Args:
        text: Text to generate embedding for

    Returns:
        List of 1024 floats (embedding vector) or None if failed
    """
    if not voyage_client:
        logger.warning("Voyage AI not configured - embeddings disabled")
        return None

    if not text or len(text.strip()) == 0:
        return None

    try:
        truncated_text = text[:30000] if len(text) > 30000 else text
        result = voyage_client.embed([truncated_text], model=EMBEDDING_MODEL, input_type="document")
        embedding = result.embeddings[0]
        logger.debug(f"Generated embedding with {len(embedding)} dimensions (Voyage {EMBEDDING_MODEL})")
        return embedding

    except Exception as e:
        logger.error(f"Embedding generation failed: {e}")
        return None


def generate_embeddings_batch(texts: List[str]) -> List[Optional[List[float]]]:
    """
    Generate embeddings for multiple texts in a single API call (much faster).

    Args:
        texts: List of texts to generate embeddings for

    Returns:
        List of embedding vectors (or None for failed items)
    """
    if not voyage_client:
        logger.warning("⚠️  Voyage AI not configured - embeddings disabled")
        return [None] * len(texts)

    if not texts:
        return []

    try:
        valid_texts = []
        valid_indices = []
        for i, text in enumerate(texts):
            if text and len(text.strip()) > 0:
                truncated = text[:30000] if len(text) > 30000 else text
                valid_texts.append(truncated)
                valid_indices.append(i)

        if not valid_texts:
            return [None] * len(texts)

        logger.info(f"🔢 Generating {len(valid_texts)} embeddings in batch (Voyage AI)...")

        # Voyage AI supports batch embedding natively
        result = voyage_client.embed(valid_texts, model=EMBEDDING_MODEL, input_type="document")

        results = [None] * len(texts)
        for i, embedding in enumerate(result.embeddings):
            results[valid_indices[i]] = embedding

        logger.info(f"✅ Generated {len(valid_texts)} embeddings successfully")
        return results

    except Exception as e:
        logger.error(f"❌ Batch embedding generation failed: {e}", exc_info=True)
        return [None] * len(texts)


def generate_query_embedding(query: str) -> Optional[List[float]]:
    """
    Generate embedding for a search query using Voyage AI voyage-3-large.

    Args:
        query: Search query text

    Returns:
        List of 1024 floats (embedding vector) or None if failed
    """
    if not voyage_client:
        logger.warning("Voyage AI not configured - embeddings disabled")
        return None

    if not query or len(query.strip()) == 0:
        return None

    try:
        result = voyage_client.embed([query], model=EMBEDDING_MODEL, input_type="query")
        return result.embeddings[0]

    except Exception as e:
        logger.error(f"Query embedding generation failed: {e}")
        return None


def analyze_file_comprehensive(text: str, filename: str = "") -> Dict:
    """Generate comprehensive detailed analysis of file content - creates detailed analysis document"""
    try:
        if not deepseek_client:
            return {
                "summary": "Analysis unavailable - API not configured.",
                "model_used": None,
                "error": "API key not configured"
            }
        
        if not text or len(text.strip()) == 0:
            return {
                "summary": "No text content to analyze.",
                "model_used": None
            }
        
        # Use much more content for ultra-detailed analysis
        analysis_text = text[:30000] if len(text) > 30000 else text
        file_hint = f" ({filename})" if filename else ""
        file_ext = os.path.splitext(filename)[1].lower() if filename else ""

        # Count images in text
        image_count = analysis_text.count("[IMAGE ")

        # Detect file category for specialized prompts
        design_exts = {'.fig', '.sketch', '.psd', '.ai', '.xd'}
        is_design = file_ext in design_exts or any(kw in analysis_text.lower() for kw in ['landing page', 'hero section', 'cta', 'navbar', 'footer', 'ui design', 'wireframe', 'mockup'])
        is_research = any(kw in analysis_text.lower() for kw in ['abstract', 'methodology', 'conclusion', 'references', 'hypothesis', 'findings', 'literature review'])

        design_section = """
DESIGN & UI ANALYSIS:
- Color palette: List ALL colors with hex codes
- Typography: All fonts, sizes, weights used
- Layout structure: Grid system, columns, spacing, responsive hints
- UI Framework: Identify if Bootstrap, Tailwind, Material UI, or custom
- Navigation: Structure, menu items, links
- CTA (Call-to-Action): Button text, placement, colors, urgency
- Hero section: Headline, subheadline, imagery, value proposition
- Visual hierarchy: What draws attention first, second, third
- Brand elements: Logo, tagline, brand colors, tone
- Conversion strategy: How the design drives user action""" if is_design else ""

        research_section = """
RESEARCH ANALYSIS:
- Research question/hypothesis
- Methodology used
- Key findings with data points
- Statistical results (p-values, confidence intervals, sample sizes)
- Limitations mentioned
- Conclusions and implications
- Citation count and key references""" if is_research else ""

        image_section = f"""
IMAGE ANALYSIS ({image_count} images detected):
- For EACH image found in the document, provide:
  * Image number and page location
  * What type of image (chart, photo, diagram, screenshot, etc.)
  * Complete description of what it shows
  * All data points, labels, and values if it's a chart/graph
  * Colors, layout, and design elements
  * How it relates to the surrounding text""" if image_count > 0 else ""

        prompt = f"""You are a document analysis expert. Analyze this document with MAXIMUM detail. Your analysis must contain MORE information than the original document by being exhaustive and organized.

Document{file_hint} ({len(analysis_text)} characters, {len(analysis_text.split())} words):
{analysis_text}

Create an EXHAUSTIVE analysis covering ALL of the following:

DOCUMENT OVERVIEW:
- Document type and category
- Main purpose (why does this document exist?)
- Target audience
- Date/time references found
- Author/source information if available

COMPLETE CONTENT BREAKDOWN:
- Section-by-section summary (cover EVERY section, not just main ones)
- Key arguments or points made in each section
- All headings and sub-headings listed

ALL DATA & FACTS:
- Every number, statistic, percentage, date, price, measurement
- Names of people, companies, products, places mentioned
- URLs, email addresses, phone numbers if present
- Technical specifications, versions, configurations
{image_section}
{design_section}
{research_section}
KEY RELATIONSHIPS & PATTERNS:
- How different parts of the document connect
- Cause-effect relationships described
- Comparisons made
- Trends or patterns in data

TERMINOLOGY & DEFINITIONS:
- Technical terms used and their meaning in context
- Acronyms and abbreviations

COMPREHENSIVE SUMMARY:
- 10-15 sentence summary covering all major points
- Nothing should be left out - if someone reads only this analysis, they should know EVERYTHING in the document

IMPORTANT: Be exhaustive. Include every detail, every number, every name. Your analysis should make the original document unnecessary to read."""

        response = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=4000
        )
        
        detailed_analysis = response.choices[0].message.content
        
        return {
            "summary": detailed_analysis,
            "model_used": "deepseek-chat"
        }
    except Exception as e:
        return {
            "summary": f"Comprehensive analysis generation failed: {str(e)}",
            "model_used": None,
            "error": str(e)
        }


def generate_page_summary(page_data: Dict, total_pages: int, filename: str = "") -> str:
    """
    Generate a detailed spatial summary for a single page.
    Covers all text, every image with exact position, all data.
    Used by generate_spatial_summary() in parallel.
    """
    if not deepseek_client:
        return ""

    page_num = page_data["page_num"]
    page_text = page_data.get("text", "").strip()
    images = page_data.get("images", [])
    heading = page_data.get("heading")  # DOCX sections have headings, PDF pages don't
    section_label = f'SECTION {page_num}: "{heading}"' if heading else f"PAGE {page_num} of {total_pages}"

    if not page_text and not images:
        return f"=== {section_label} ===\n(Empty)\n"

    # Build image section for prompt
    images_block = ""
    if images:
        images_block = f"\n\nIMAGES IN THIS SECTION ({len(images)} images):\n"
        for i, img in enumerate(images):
            images_block += f"\n[IMAGE {i + 1}]:\n{img['text']}\n"

    prompt = f"""You are analyzing {section_label} from the document "{filename}".

PAGE TEXT:
{page_text}
{images_block}

Create an EXHAUSTIVE SPATIAL SUMMARY of this page. You MUST cover ALL of the following:

LAYOUT & STRUCTURE:
- Describe what is placed at: top, bottom, left, right, center, full-width
- Column structure (1 column, 2 columns, grid, etc.)
- Any headers, footers, sidebars

TEXT CONTENT:
- Every heading and subheading (exact text)
- All paragraph content summarized (nothing skipped)
- Any captions, labels, callouts, tooltips
- Bullet points or lists (all items)
- Any highlighted or bold text

IMAGES (MANDATORY - do NOT skip any image):
For EACH image listed above:
- Exact position on page (e.g. top-right, center-left, bottom-full-width)
- Image type (photo, chart, bar graph, pie chart, logo, icon, diagram, screenshot, etc.)
- Full description of what it shows
- ALL text, numbers, labels visible in the image
- Colors, design style
- How it connects to the surrounding text
- If chart/graph: all axis labels, all data values, trends

DATA & NUMBERS:
- Every number, price, percentage, date, measurement on this page
- Any statistics or metrics

KEY POINTS:
- The 3-5 most important things on this page

CONTINUITY (IMPORTANT):
- If content on this page clearly continues from a previous page, start with: "CONTINUED FROM PAGE {page_num - 1}:"
- If content on this page continues onto the next page, end with: "CONTINUES ON PAGE {page_num + 1}"
- Always note page number references so everything stays aligned in the full document

IMPORTANT: Be completely exhaustive. If a user later asks "what is the image on the right of page {page_num}" or "what does the chart on page {page_num} show" — your summary must answer it fully."""

    try:
        response = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=2000
        )
        page_summary = response.choices[0].message.content
        return f"{'='*60}\n=== {section_label} ===\n{'='*60}\n{page_summary}\n"
    except Exception as e:
        logger.error(f"{section_label} summary failed: {e}")
        return f"=== {section_label} ===\n(Summary failed: {e})\n"


def generate_spatial_summary(text_data: Dict, filename: str = "") -> Dict:
    """
    Generate a rich spatial summary of a document.

    For PDFs with per-page data: runs generate_page_summary() for ALL pages
    in parallel (one DeepSeek call per page), then merges into master summary.

    For other file types: single DeepSeek call with strict instructions.

    Returns:
        Dict with 'summary' (str) and 'model_used' (str)
    """
    if not deepseek_client:
        return {"summary": "Analysis unavailable - DeepSeek API not configured.", "model_used": None}

    pages = text_data.get("metadata", {}).get("pages")
    full_text = text_data.get("text", "")

    # PDF with per-page structure → parallel per-page summaries
    if pages and len(pages) > 0:
        total_pages = len(pages)
        logger.info(f"📄 Generating spatial summary: {total_pages} pages in parallel for '{filename}'")

        page_summaries = [""] * total_pages

        def _summarize_page(page_data):
            idx = page_data["page_num"] - 1
            result = generate_page_summary(page_data, total_pages, filename)
            return idx, result

        with ThreadPoolExecutor(max_workers=min(total_pages, 20)) as executor:
            futures = {executor.submit(_summarize_page, p): p for p in pages}
            for future in as_completed(futures):
                try:
                    idx, result = future.result()
                    page_summaries[idx] = result
                    logger.info(f"  ✅ Page {idx + 1}/{total_pages} summary done")
                except Exception as e:
                    logger.warning(f"  ⚠️  Page summary failed: {e}")

        # Merge all page summaries into master summary — fully aligned 1 to N
        master = (
            f"{'='*70}\n"
            f"DOCUMENT: {filename}\n"
            f"TOTAL PAGES: {total_pages}\n"
            f"{'='*70}\n\n"
        )
        for i, ps in enumerate(page_summaries):
            master += ps if ps else f"{'='*60}\n=== PAGE {i+1} of {total_pages} ===\n{'='*60}\n(No content)\n"
            master += "\n"

        logger.info(f"✅ Spatial summary complete: {len(master)} chars across {total_pages} pages")
        return {"summary": master, "model_used": "deepseek-chat"}

    # Non-PDF or no page structure → single comprehensive call
    if not full_text or len(full_text.strip()) == 0:
        return {"summary": "No text content to analyze.", "model_used": None}

    logger.info(f"📝 Generating spatial summary (single call) for '{filename}'")
    analysis_text = full_text[:60000] if len(full_text) > 60000 else full_text
    image_count = analysis_text.count("[IMAGE ")

    image_instruction = ""
    if image_count > 0:
        image_instruction = f"""
IMAGES ({image_count} images detected — MANDATORY):
For EACH [IMAGE ...] block found in the text:
- Image number and exact location in document
- Image type (photo, chart, diagram, logo, etc.)
- Full detailed description
- ALL text, numbers, labels in the image
- Colors and design
- Position relative to surrounding text (left, right, top, bottom, inline)
- If chart/graph: ALL axis labels and data values"""

    prompt = f"""You are a document analysis expert. Analyze this document with MAXIMUM spatial and visual detail.

Document: "{filename}" ({len(analysis_text)} characters)
{analysis_text}

Create an EXHAUSTIVE SPATIAL ANALYSIS covering:

DOCUMENT OVERVIEW:
- Document type, purpose, target audience
- Overall layout and structure

SECTION-BY-SECTION BREAKDOWN:
- Every section and subsection (nothing skipped)
- All content, data, facts in each section
{image_instruction}

ALL DATA & FACTS:
- Every number, price, date, percentage, measurement
- All names, companies, products, URLs

KEY RELATIONSHIPS:
- How sections and visuals connect
- Important patterns and insights

COMPREHENSIVE SUMMARY:
- 10-15 sentences covering everything

IMPORTANT: Be exhaustive. Every image must be described with its position and full content."""

    try:
        response = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=4000
        )
        return {"summary": response.choices[0].message.content, "model_used": "deepseek-chat"}
    except Exception as e:
        logger.error(f"Spatial summary failed: {e}")
        return {"summary": f"Summary generation failed: {str(e)}", "model_used": None, "error": str(e)}


async def enrich_summary_with_web_search(summary: str, filename: str = "") -> str:
    """Enrich a summary with web search results for additional context"""
    try:
        from app.services.web_search import search_web

        if not summary or len(summary) < 100:
            return summary

        # Extract key topics from summary for web search (first 500 chars)
        topic_text = summary[:500]

        # Build search query from filename and key content
        search_query = filename.replace('.', ' ').replace('_', ' ').replace('-', ' ')
        # Add first meaningful sentence
        first_line = topic_text.split('\n')[0][:100] if topic_text else ""
        if first_line:
            search_query = f"{search_query} {first_line}"

        search_query = search_query.strip()[:100]  # Limit query length

        if not search_query or len(search_query) < 5:
            return summary

        logger.info(f"🌐 Enriching summary with web search: '{search_query[:50]}...'")

        results = await search_web(search_query, num_results=3)

        if results:
            web_context = "\n\nADDITIONAL CONTEXT FROM WEB:\n"
            for r in results:
                web_context += f"- {r.get('title', '')}: {r.get('snippet', '')} (Source: {r.get('displayLink', '')})\n"

            return summary + web_context

        return summary
    except Exception as e:
        logger.warning(f"Web enrichment failed (non-critical): {e}")
        return summary


def generate_summary(text: str, filename: str = "", model_name: str = "deepseek-chat") -> Dict:
    """Generate file summary - uses comprehensive analysis for detailed summaries"""
    # Use comprehensive analysis for better summaries
    result = analyze_file_comprehensive(text, filename)
    return {
        "summary": result.get("summary", "Summary generation failed."),
        "model_used": result.get("model_used", model_name),
        "error": result.get("error")
    }


def fetch_full_file_content(storage_path: str, supabase) -> str:
    """
    Fetch and extract complete content from original file in Supabase Storage.
    
    Args:
        storage_path: Path to file in Supabase Storage (e.g., 'user_id/bucket_id/filename')
        supabase: Supabase client instance
    
    Returns:
        str: Full extracted text content from the original file
    """
    logger.info(f"📥 Fetching full file from storage: {storage_path}")
    
    try:
        # Download file from Supabase Storage
        logger.debug(f"  🔽 Downloading from storage...")
        file_data = supabase.storage.from_("files").download(storage_path)
        
        if not file_data:
            logger.error(f"❌ File not found in storage: {storage_path}")
            return ""
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(storage_path)[1]) as temp_file:
            temp_file.write(file_data)
            temp_path = temp_file.name
        
        logger.debug(f"  📄 Saved to temp: {temp_path}")
        
        # Detect MIME type from extension
        file_ext = os.path.splitext(storage_path)[1].lower()
        mime_type_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.webp': 'image/webp'
        }
        mime_type = mime_type_map.get(file_ext, 'application/octet-stream')
        
        # Extract text using existing function
        logger.info(f"  🔍 Extracting text (MIME: {mime_type})...")
        text_data = extract_text_from_file(temp_path, mime_type)
        full_text = text_data["text"]
        
        # Clean up temp file
        os.unlink(temp_path)
        
        logger.info(f"  ✅ Full file content extracted: {len(full_text)} chars")
        return full_text
        
    except Exception as e:
        logger.error(f"❌ Error fetching full file: {str(e)}")
        return ""
